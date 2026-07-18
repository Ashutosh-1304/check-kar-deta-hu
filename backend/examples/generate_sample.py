import os
import sys

# Ensure backend directory is in path when running directly
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx import Document
from docx.shared import Inches, Pt

def generate_sample():
    doc = Document()
    
    # Metadata
    core_props = doc.core_properties
    core_props.title = "Sample Compliance Document"
    core_props.author = "Test Author"
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Heading
    doc.add_heading('Main Title', level=1)
    
    # Paragraph
    p = doc.add_paragraph('This is a test paragraph with ')
    run = p.add_run('bold text')
    run.bold = True
    p.add_run(' and ')
    run2 = p.add_run('italic text')
    run2.italic = True
    p.add_run('.')
    
    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Row 1, Col 1'
    table.cell(0, 1).text = 'Row 1, Col 2'
    table.cell(1, 0).text = 'Row 2, Col 1'
    table.cell(1, 1).text = 'Row 2, Col 2'
    
    examples_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'examples')
    os.makedirs(examples_dir, exist_ok=True)
    doc.save(os.path.join(examples_dir, 'sample_document.docx'))
    print("Sample document generated successfully.")

if __name__ == "__main__":
    generate_sample()
