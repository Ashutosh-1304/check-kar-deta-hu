from typing import Union
from pathlib import Path
from io import BytesIO
import docx
from app.parser.base import BaseParser
from app.parser.models import (
    DocumentModel, DocumentMetadata, SectionModel, PageProperties,
    ParagraphModel, HeadingModel, RunModel, TableModel, ImageModel
)

class DocxParser(BaseParser):
    def parse(self, file: Union[str, Path, BytesIO]) -> DocumentModel:
        doc = docx.Document(file)
        
        # 1. Extract Metadata
        core_props = doc.core_properties
        metadata = DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            core_properties={
                "subject": core_props.subject,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
            }
        )
        
        sections_models = []
        
        # 2. Extract Sections (Page properties)
        for section in doc.sections:
            page_props = PageProperties(
                page_width=section.page_width.inches if section.page_width else None,
                page_height=section.page_height.inches if section.page_height else None,
                margin_top=section.top_margin.inches if section.top_margin else None,
                margin_bottom=section.bottom_margin.inches if section.bottom_margin else None,
                margin_left=section.left_margin.inches if section.left_margin else None,
                margin_right=section.right_margin.inches if section.right_margin else None
            )
            
            section_model = SectionModel(
                page_properties=page_props,
                paragraphs=[],
                tables=[],
                images=[]
            )
            sections_models.append(section_model)
            
        if not sections_models:
            sections_models.append(SectionModel(page_properties=PageProperties()))
            
        primary_section = sections_models[0]
        
        # 3. Extract Paragraphs & Headings
        for para in doc.paragraphs:
            runs_models = []
            for run in para.runs:
                font = run.font
                color = None
                if font.color and font.color.rgb:
                    color = str(font.color.rgb)
                
                runs_models.append(RunModel(
                    text=run.text,
                    font_family=font.name,
                    font_size=font.size.pt if font.size else None,
                    bold=font.bold,
                    italic=font.italic,
                    underline=font.underline,
                    color=color
                ))
            
            style_name = para.style.name if para.style else None
            
            # handle alignment which might be an enum object
            alignment = None
            if para.alignment is not None:
                try:
                    alignment = para.alignment.name
                except AttributeError:
                    alignment = str(para.alignment)
            
            para_format = para.paragraph_format
            first_line_indent = para_format.first_line_indent.inches if para_format and para_format.first_line_indent else None
            space_after = para_format.space_after.pt if para_format and para_format.space_after else None

            # Detect Headings
            if style_name and style_name.startswith('Heading'):
                try:
                    level = int(style_name.split(' ')[1])
                except (ValueError, IndexError):
                    level = 1
                
                primary_section.paragraphs.append(HeadingModel(
                    text=para.text,
                    runs=runs_models,
                    style=style_name,
                    alignment=alignment,
                    first_line_indent=first_line_indent,
                    space_after=space_after,
                    level=level
                ))
            else:
                primary_section.paragraphs.append(ParagraphModel(
                    text=para.text,
                    runs=runs_models,
                    style=style_name,
                    alignment=alignment,
                    first_line_indent=first_line_indent,
                    space_after=space_after
                ))
                
        # 4. Extract Tables
        for table in doc.tables:
            primary_section.tables.append(TableModel(
                rows=len(table.rows),
                cols=len(table.columns),
                style=table.style.name if table.style else None
            ))
            
        # 5. Extract Images
        for inline_shape in doc.inline_shapes:
            primary_section.images.append(ImageModel(
                width=inline_shape.width.inches if inline_shape.width else None,
                height=inline_shape.height.inches if inline_shape.height else None
            ))
            
        return DocumentModel(
            metadata=metadata,
            sections=sections_models
        )
